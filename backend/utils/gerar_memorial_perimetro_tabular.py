r"""Gera um memorial tabular de perímetro com quantidade dinâmica de trechos.

Dependências:
    pip install python-docx Pillow

Exemplo de uso no terminal do VS Code (PowerShell):
    python gerar_memorial_perimetro_tabular.py `
        --tabular "C:\caminho\MD_Perímetro_Gutierrez_Tabular.docx" `
        --descritivo "C:\caminho\MD_Perímetro_Loteamento_gutierrez.docx" `
        --saida "C:\caminho\MD_Perímetro_Gutierrez_Atualizado.docx"

O script preserva os arquivos de entrada e cria um novo DOCX.
"""

from __future__ import annotations

import argparse
import copy
import io
import math
import re
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


CANVAS_WIDTH = 2000
CANVAS_HEIGHT = 1488
BLACK = "#050505"
BROWN = "#9A2E1F"
LEADER = "#737373"


@dataclass
class PerimeterSegment:
    source: str
    destination: str
    easting: str
    northing: str
    confrontation: str
    distance: str


@dataclass
class PerimeterMemorial:
    property_name: str
    owner: str
    municipality: str
    area: str
    perimeter: str
    technician: str
    crea: str
    art: str
    segments: list[PerimeterSegment]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Atualiza o memorial tabular e recria o croqui do perímetro."
    )
    parser.add_argument("--tabular", required=True, type=Path, help="DOCX tabular utilizado como modelo.")
    parser.add_argument("--descritivo", required=True, type=Path, help="DOCX com a descrição do perímetro.")
    parser.add_argument("--saida", type=Path, help="DOCX de saída. Se omitido, usa o nome *_Atualizado.docx.")
    return parser.parse_args()


def normalize_vertex(value: str) -> str:
    value = re.sub(r"\s+", "", value.strip().strip(" ,;:()[]{}"))
    value = value.replace("–", "-").replace("—", "-").upper()
    simple = re.fullmatch(r"([A-ZÀ-Ý]+)-?(\d+)", value)
    if simple:
        return f"{simple.group(1)}-{int(simple.group(2)):02d}"
    return value


def normalize_vertex_in_sequence(value: str, previous: str = "") -> str:
    vertex = normalize_vertex(value)
    if vertex.isdigit() and previous:
        previous_match = re.fullmatch(r"([A-ZÀ-Ý]+)-(\d+)", previous)
        if previous_match:
            return f"{previous_match.group(1)}-{int(vertex):02d}"
    return vertex


def parse_pt_number(value: str) -> float:
    return float(value.strip().replace(".", "").replace(",", "."))


def format_m(value: float) -> str:
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + " m"


def set_cell_text_preserve(cell, text: str, reference_cell=None) -> None:
    paragraph = cell.paragraphs[0]
    target_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    for run in paragraph.runs:
        run.text = ""

    if reference_cell is not None:
        ref_paragraph = reference_cell.paragraphs[0]
        if ref_paragraph._p.pPr is not None:
            if paragraph._p.pPr is not None:
                paragraph._p.remove(paragraph._p.pPr)
            paragraph._p.insert(0, copy.deepcopy(ref_paragraph._p.pPr))
        ref_run = next((run for run in ref_paragraph.runs if run.text.strip()), None)
        if ref_run is not None and ref_run._r.rPr is not None:
            if target_run._r.rPr is not None:
                target_run._r.remove(target_run._r.rPr)
            target_run._r.insert(0, copy.deepcopy(ref_run._r.rPr))

    target_run.text = text


def field_value(paragraphs: list[str], label: str) -> str:
    for text in paragraphs:
        match = re.match(rf"\s*{re.escape(label)}\s*:\s*(.*?)\s*$", text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return ""


def first_field_value(paragraphs: list[str], *labels: str) -> str:
    for label in labels:
        value = field_value(paragraphs, label)
        if value:
            return value
    return ""


def document_text_blocks(document: Document) -> list[str]:
    """Lê parágrafos do corpo e de tabelas sem modificar o DOCX de origem."""
    blocks: list[str] = []
    for paragraph in document.element.body.xpath(".//w:p"):
        text = "".join(node.text or "" for node in paragraph.xpath(".//w:t")).strip()
        if text:
            blocks.append(text)
    return blocks


VERTEX_TOKEN = r"[A-ZÀ-Ý0-9]+(?:\s*[-–—_.]\s*[A-ZÀ-Ý0-9]+)*"


def extract_coordinates(description: str) -> tuple[dict[str, tuple[str, str]], list[str]]:
    coordinates: dict[str, tuple[str, str]] = {}
    order: list[str] = []
    number = r"[0-9][0-9.,]*"
    prefix = rf"v[ée]rtice\s+(?P<vertex>{VERTEX_TOKEN})\s*,?\s*(?:de\s+)?coordenadas?\s+"
    patterns = (
        re.compile(
            prefix
            + rf"N\s*[:=]?\s*(?P<north>{number})\s*m?\s*(?:e|,)\s*"
            + rf"E\s*[:=]?\s*(?P<east>{number})\s*m?",
            re.IGNORECASE,
        ),
        re.compile(
            prefix
            + rf"E\s*[:=]?\s*(?P<east>{number})\s*m?\s*(?:e|,)\s*"
            + rf"N\s*[:=]?\s*(?P<north>{number})\s*m?",
            re.IGNORECASE,
        ),
        re.compile(
            prefix
            + rf"(?:EX|X)\s*[:=]?\s*(?P<east>{number})\s*[,;]?\s*"
            + rf"(?:NY|Y)\s*[:=]?\s*(?P<north>{number})",
            re.IGNORECASE,
        ),
    )
    matches = sorted(
        (match for pattern in patterns for match in pattern.finditer(description)),
        key=lambda match: match.start(),
    )
    for match in matches:
        vertex = normalize_vertex_in_sequence(match.group("vertex"), order[-1] if order else "")
        coordinate = (match.group("east"), match.group("north"))
        if vertex in coordinates and coordinates[vertex] != coordinate and order:
            previous_match = re.fullmatch(r"([A-ZÀ-Ý]+)-(\d+)", order[-1])
            duplicate_match = re.fullmatch(r"([A-ZÀ-Ý]+)-(\d+)", vertex)
            if previous_match and duplicate_match and previous_match.group(1) == duplicate_match.group(1):
                inferred = f"{previous_match.group(1)}-{int(previous_match.group(2)) + 1:02d}"
                if inferred not in coordinates:
                    vertex = inferred
        if vertex not in coordinates:
            order.append(vertex)
        coordinates[vertex] = coordinate
    return coordinates, order


def confrontation_from_interval(interval: str, current: str) -> str:
    patterns = (
        r"(?:segue|prossegue|continua)?\s*(?:confrontando|limitando-se)\s+com\s+",
        r"situad[oa]\s+(?:em\s+comum\s+com|no\s+limite\s+d[ao])\s+",
    )
    candidates: list[tuple[int, str]] = []
    for pattern in patterns:
        for match in re.finditer(pattern, interval, re.IGNORECASE):
            candidates.append((match.end(), interval[match.end() :]))
    if not candidates:
        return current
    _, value = max(candidates, key=lambda item: item[0])
    value = re.split(
        r",\s*(?:com\s+os\s+seguintes\s+azimutes?\s+e\s+dist[âa]ncias?\s*:|"
        r"com\s+azimute|azimute)|;\s*(?:deste|da[ií])\b",
        value,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0]
    return " ".join(value.strip(" ,;:").split()) or current


def format_pt_number(value: float) -> str:
    return f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def extract_perimeter_memorial(descriptive_path: Path) -> PerimeterMemorial:
    source = Document(descriptive_path)
    paragraphs = document_text_blocks(source)
    description_blocks = [
        text
        for text in paragraphs
        if re.search(r"\b(?:azimute|azimutes)\b", text, re.IGNORECASE)
        and re.search(r"\bcoordenadas?\b", text, re.IGNORECASE)
    ]
    description = " ".join(description_blocks)
    if not description:
        raise ValueError(
            "A descrição técnica com coordenadas, azimutes e distâncias não foi encontrada no DOCX."
        )

    coordinates, coordinate_order = extract_coordinates(description)
    if not coordinate_order:
        raise ValueError(
            "Nenhuma coordenada de vértice foi reconhecida. São aceitos N/E, E/N, NY/EX e Y/X."
        )

    dms_angle = r"\d{1,3}\s*[°º]\s*\d{1,2}\s*['’´]\s*\d{1,2}(?:[.,]\d+)?\s*(?:[\"”″]|'{2}|’{2})"
    segment_pattern = re.compile(
        rf"(?:(?:com\s+)?azimute(?:\s+de)?\s+)?{dms_angle}\s*(?:e|,)\s*"
        rf"(?:(?:com\s+)?(?:uma\s+)?dist[âa]ncia(?:\s+de)?\s+)?"
        rf"(?P<distance>[0-9][0-9.,]*)\s*m\s*at[ée]\s+(?:o\s+)?v[ée]rtice\s+"
        rf"(?P<destination>{VERTEX_TOKEN})",
        re.IGNORECASE,
    )

    segments: list[PerimeterSegment] = []
    current = coordinate_order[0]
    confrontation = ""
    previous_end = 0
    for index, match in enumerate(segment_pattern.finditer(description), start=1):
        interval = description[previous_end : match.start()]
        confrontation = confrontation_from_interval(interval, confrontation)
        destination = normalize_vertex_in_sequence(match.group("destination"), current)
        if destination == segments[0].source if segments else destination == current:
            remaining_text = description[match.end() :]
            if re.search(rf"{dms_angle}\s*(?:e|,)", remaining_text, re.IGNORECASE):
                current_match = re.fullmatch(r"([A-ZÀ-Ý]+)-(\d+)", current)
                if current_match:
                    inferred = f"{current_match.group(1)}-{int(current_match.group(2)) + 1:02d}"
                    if inferred in coordinates:
                        destination = inferred
        if current not in coordinates:
            raise ValueError(f"Coordenada ausente para {current} no trecho {index}.")
        if not confrontation:
            raise ValueError(f"Confrontante ausente no trecho {current}–{destination}.")
        easting, northing = coordinates[current]
        segments.append(
            PerimeterSegment(
                source=current,
                destination=destination,
                easting=easting,
                northing=northing,
                confrontation=confrontation,
                distance=match.group("distance"),
            )
        )
        current = destination
        previous_end = match.end()

    if not segments:
        raise ValueError(
            "Nenhum trecho foi reconhecido. Verifique se cada trecho contém azimute em graus, "
            "distância em metros e o vértice de destino."
        )
    if segments[-1].destination != segments[0].source:
        raise ValueError(
            f"O perímetro não fecha: termina em {segments[-1].destination} e começa em {segments[0].source}."
        )
    # Um vértice pode aparecer mais de uma vez em perímetros compostos ou com ramificações.
    # Por isso, a quantidade de trechos não é obrigada a ser igual à de coordenadas únicas.

    raw_area = first_field_value(paragraphs, "Área", "Área total", "Área do imóvel")
    area_match = re.search(r"[\d.,]+", raw_area)
    if not area_match:
        raise ValueError("A área não foi encontrada no memorial descritivo.")
    area = area_match.group(0)
    if "ha" in raw_area.casefold():
        area = format_pt_number(parse_pt_number(area) * 10_000)

    raw_perimeter = first_field_value(paragraphs, "Perímetro", "Perímetro total")
    perimeter_match = re.search(r"[\d.,]+", raw_perimeter)
    if not perimeter_match:
        raise ValueError("O perímetro não foi encontrado no memorial descritivo.")
    perimeter = perimeter_match.group(0)
    # O gerador transpõe os valores do memorial sem substituir o papel de uma
    # conferência topográfica. Divergências existentes no original não bloqueiam a saída.

    crea = field_value(paragraphs, "CREA")
    art = field_value(paragraphs, "ART")
    technician = ""
    if crea:
        crea_index = next(
            (index for index, text in enumerate(paragraphs) if re.match(r"\s*CREA\s*:", text, re.IGNORECASE)),
            None,
        )
        if crea_index is not None:
            technician = next(
                (
                    text
                    for text in reversed(paragraphs[:crea_index])
                    if text.strip() and not re.fullmatch(r"_+", text.strip())
                ),
                "",
            )

    return PerimeterMemorial(
        property_name=first_field_value(paragraphs, "Imóvel", "Propriedade", "Loteamento")
        or descriptive_path.stem,
        owner=first_field_value(paragraphs, "Proprietário", "Interessado", "Requerente"),
        municipality=first_field_value(paragraphs, "Município", "Município/UF", "Cidade"),
        area=area,
        perimeter=perimeter,
        technician=technician,
        crea=crea,
        art=art,
        segments=segments,
    )


def find_vertex_table(doc: Document):
    for table in doc.tables:
        if not table.rows or len(table.columns) < 7:
            continue
        header = [cell.text.strip().casefold() for cell in table.rows[0].cells]
        if header[0] == "de" and header[1] == "para":
            return table
    raise ValueError("A tabela de vértices (De/Para) não foi encontrada no memorial tabular.")


def update_table(doc: Document, memorial: PerimeterMemorial):
    table = find_vertex_table(doc)
    set_cell_text_preserve(table.cell(0, 2), "Coordenada E")
    set_cell_text_preserve(table.cell(0, 3), "Coordenada N")
    set_cell_text_preserve(table.cell(0, 6), "Complemento")
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    if not header_properties.xpath("./w:tblHeader"):
        repeat_header = OxmlElement("w:tblHeader")
        repeat_header.set(qn("w:val"), "true")
        header_properties.append(repeat_header)

    row_model = copy.deepcopy(table.rows[1]._tr)
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for segment in memorial.segments:
        table._tbl.append(copy.deepcopy(row_model))
        row = table.rows[-1]
        values = (
            segment.source,
            segment.destination,
            segment.easting,
            segment.northing,
            segment.confrontation,
            segment.distance,
            "",
        )
        for column, value in enumerate(values):
            reference = row.cells[5] if column == 6 else None
            set_cell_text_preserve(row.cells[column], value, reference_cell=reference)

    return table


def update_metadata(doc: Document, memorial: PerimeterMemorial) -> None:
    title = f"MEMORIAL TABULAR DO PERÍMETRO DO {memorial.property_name.upper()}"
    paragraph = doc.paragraphs[0]
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    for item in paragraph.runs:
        item.text = ""
    run.text = title

    info = doc.tables[0]
    values = (
        ("Código da Parcela:", ""),
        (f"Imóvel: {memorial.property_name}", "Cartório (CNS):"),
        (f"Proprietário: {memorial.owner}", "CNPJ nº:"),
        ("Endereço:", f"Município/UF: {memorial.municipality}"),
        ("Sistema Geodésico de Referência: SIRGAS2000", "Projeção cartográfica de distância e área: U T M"),
        ("Modelo de conversão de altitudes", ""),
        (f"Área (m²): {memorial.area} m²", f"Perímetro (m): {memorial.perimeter} m"),
    )
    for row, (left, right) in zip(info.rows, values):
        set_cell_text_preserve(row.cells[0], left)
        set_cell_text_preserve(row.cells[1], right)

    signature = next(
        (paragraph for paragraph in doc.paragraphs if "CREA:" in paragraph.text and "ART:" in paragraph.text),
        None,
    )
    if signature is not None and (memorial.technician or memorial.crea or memorial.art):
        technician = memorial.technician or "Responsável técnico"
        text = (
            "_______________________________________________\n"
            f"{technician}\n"
            f"CREA: {memorial.crea}\n"
            f"ART: {memorial.art}"
        )
        run = signature.runs[0] if signature.runs else signature.add_run()
        for item in signature.runs:
            item.text = ""
        run.text = text


def records_from_table(table) -> list[dict]:
    records = []
    for row in table.rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        records.append(
            {
                "de": normalize_vertex(values[0]),
                "para": normalize_vertex(values[1]),
                "e": parse_pt_number(values[2]),
                "n": parse_pt_number(values[3]),
                "confrontante": values[4],
                "distancia": parse_pt_number(values[5]),
            }
        )
    if records[-1]["para"] != records[0]["de"]:
        raise ValueError("A tabela não fecha o perímetro no vértice inicial.")
    return records


def load_font(filename: str, size: int, fallback: str = "arial.ttf"):
    fonts = Path(r"C:\Windows\Fonts")
    for candidate in (fonts / filename, fonts / fallback):
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def wrap_text(draw, text: str, font, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for word in text.split():
        trial = word if not current else current + " " + word
        if draw.textbbox((0, 0), trial, font=font)[2] <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def overlap_area(a, b) -> float:
    x0 = max(a[0], b[0])
    y0 = max(a[1], b[1])
    x1 = min(a[2], b[2])
    y1 = min(a[3], b[3])
    return max(0, x1 - x0) * max(0, y1 - y0)


def spread(values: list[float], low: float, high: float, gap: float) -> list[float]:
    if not values:
        return []
    values = sorted(values)
    result = [max(low, values[0])]
    for value in values[1:]:
        result.append(max(value, result[-1] + gap))
    if result[-1] > high:
        shift = result[-1] - high
        result = [value - shift for value in result]
    if result[0] < low:
        shift = low - result[0]
        result = [value + shift for value in result]
    return result


def group_boundaries(records: list[dict]) -> list[dict]:
    groups: list[dict] = []
    for index, record in enumerate(records):
        if groups and groups[-1]["name"] == record["confrontante"]:
            groups[-1]["indices"].append(index)
            groups[-1]["distance"] += record["distancia"]
        else:
            groups.append(
                {"name": record["confrontante"], "indices": [index], "distance": record["distancia"]}
            )
    return groups


def render_croqui(records: list[dict]) -> bytes:
    if len(records) > 80:
        vertex_size, callout_size = 15, 20
    elif len(records) > 50:
        vertex_size, callout_size = 18, 22
    elif len(records) > 35:
        vertex_size, callout_size = 20, 24
    elif len(records) > 30:
        vertex_size, callout_size = 25, 25
    elif len(records) > 20:
        vertex_size, callout_size = 31, 27
    else:
        vertex_size, callout_size = 35, 29
    vertex_font = load_font("timesbd.ttf", vertex_size)
    callout_font = load_font("times.ttf", callout_size)
    north_font = load_font("timesbd.ttf", 48)

    image = Image.new("RGB", (CANVAS_WIDTH, CANVAS_HEIGHT), "white")
    draw = ImageDraw.Draw(image)

    plot = (570, 185, 1405, 1305)
    x0, y0, x1, y1 = plot
    eastings = [record["e"] for record in records]
    northings = [record["n"] for record in records]
    emin, emax = min(eastings), max(eastings)
    nmin, nmax = min(northings), max(northings)
    scale = min((x1 - x0) / (emax - emin), (y1 - y0) / (nmax - nmin))
    used_width = (emax - emin) * scale
    used_height = (nmax - nmin) * scale
    offset_x = x0 + (x1 - x0 - used_width) / 2
    offset_y = y0 + (y1 - y0 - used_height) / 2

    def xy(easting: float, northing: float) -> tuple[float, float]:
        return (
            offset_x + (easting - emin) * scale,
            offset_y + (nmax - northing) * scale,
        )

    points = [xy(record["e"], record["n"]) for record in records]
    centroid = (
        sum(point[0] for point in points) / len(points),
        sum(point[1] for point in points) / len(points),
    )

    groups = group_boundaries(records)
    for group in groups:
        midpoints = []
        weights = []
        for index in group["indices"]:
            point_a = points[index]
            point_b = points[(index + 1) % len(points)]
            midpoints.append(((point_a[0] + point_b[0]) / 2, (point_a[1] + point_b[1]) / 2))
            weights.append(records[index]["distancia"])
        total = sum(weights)
        midpoint_x = sum(point[0] * weight for point, weight in zip(midpoints, weights)) / total
        midpoint_y = sum(point[1] * weight for point, weight in zip(midpoints, weights)) / total
        group["mid"] = (midpoint_x, midpoint_y)
        dx, dy = midpoint_x - centroid[0], midpoint_y - centroid[1]
        if abs(dx) > abs(dy) * 1.15:
            group["side"] = "right" if dx > 0 else "left"
        else:
            group["side"] = "bottom" if dy > 0 else "top"

    for side in ("left", "right"):
        selected = sorted((group for group in groups if group["side"] == side), key=lambda g: g["mid"][1])
        positions = spread([group["mid"][1] for group in selected], 245, 1235, 118)
        for group, value in zip(selected, positions):
            group["label_pos"] = (360 if side == "left" else 1470, value)

    for side in ("top", "bottom"):
        selected = sorted((group for group in groups if group["side"] == side), key=lambda g: g["mid"][0])
        if side == "top" and len(selected) > 1:
            positions = [500 + index * (1000 / (len(selected) - 1)) for index in range(len(selected))]
        elif side == "bottom" and len(selected) > 1:
            positions = [470 + index * (1000 / (len(selected) - 1)) for index in range(len(selected))]
        else:
            positions = spread([group["mid"][0] for group in selected], 530, 1440, 250)
        for index, (group, value) in enumerate(zip(selected, positions)):
            group["label_pos"] = (value, 55 if side == "top" else 1370)
            if side == "bottom":
                group["route_y"] = 1210 + index * 30

    # Linhas e textos de confrontação.
    for group in groups:
        midpoint_x, midpoint_y = group["mid"]
        label_x, label_y = group["label_pos"]
        side = group["side"]
        if side == "left":
            draw.line(
                (midpoint_x, midpoint_y, 475, midpoint_y, 475, label_y, label_x + 8, label_y),
                fill=BROWN,
                width=4,
            )
            anchor = "rm"
        elif side == "right":
            draw.line(
                (midpoint_x, midpoint_y, 1450, midpoint_y, 1450, label_y, label_x - 8, label_y),
                fill=BROWN,
                width=4,
            )
            anchor = "lm"
        elif side == "top":
            draw.line(
                (midpoint_x, midpoint_y, midpoint_x, 185, label_x, 185, label_x, 160),
                fill=BROWN,
                width=4,
            )
            anchor = "ma"
        else:
            route_y = group["route_y"]
            draw.line(
                (
                    midpoint_x,
                    midpoint_y,
                    midpoint_x,
                    route_y,
                    label_x,
                    route_y,
                    label_x,
                    label_y - 35,
                ),
                fill=BROWN,
                width=4,
            )
            anchor = "ms"

        width = 205 if side == "right" else 245
        lines = wrap_text(draw, group["name"], callout_font, width)
        text = "\n".join(lines + [format_m(group["distance"])])
        draw.multiline_text(
            (label_x, label_y),
            text,
            fill=BROWN,
            font=callout_font,
            anchor=anchor,
            align="center",
            spacing=0,
        )

    # Perímetro e marcadores.
    for index in range(len(points)):
        draw.line((*points[index], *points[(index + 1) % len(points)]), fill="white", width=16)
        draw.line((*points[index], *points[(index + 1) % len(points)]), fill=BLACK, width=10)
    for x, y in points:
        draw.ellipse((x - 12, y - 12, x + 12, y + 12), fill=BLACK)

    # Identificação automática dos vértices, com prevenção de sobreposição.
    placed = []
    standard_angles = [index * math.pi / 8 for index in range(16)]
    for (x, y), record in zip(points, records):
        label = record["de"]
        text_box = draw.textbbox((0, 0), label, font=vertex_font)
        text_width = text_box[2] - text_box[0]
        text_height = text_box[3] - text_box[1]
        radial = math.atan2(y - centroid[1], x - centroid[0])
        ordered_angles = [radial] + sorted(
            standard_angles,
            key=lambda angle: abs(math.atan2(math.sin(angle - radial), math.cos(angle - radial))),
        )
        candidates = []
        for distance in (30, 48, 70, 96, 128, 165, 210):
            for angle in ordered_angles:
                center_x = x + math.cos(angle) * distance
                center_y = y + math.sin(angle) * distance
                box = (
                    center_x - text_width / 2 - 5,
                    center_y - text_height / 2 - 3,
                    center_x + text_width / 2 + 5,
                    center_y + text_height / 2 + 3,
                )
                overlap = sum(overlap_area(box, other) for other in placed)
                outside = (
                    max(0, 500 - box[0])
                    + max(0, box[2] - 1500)
                    + max(0, 145 - box[1])
                    + max(0, box[3] - 1340)
                )
                score = overlap * 1000 + outside * 100 + distance
                candidates.append((score, box, center_x, center_y, distance))
        _, box, center_x, center_y, distance = min(candidates, key=lambda item: item[0])
        if distance >= 56:
            draw.line((x, y, center_x, center_y), fill=LEADER, width=2)
        draw.rectangle(box, fill="white")
        draw.text((center_x, center_y), label, fill=BLACK, font=vertex_font, anchor="mm")
        placed.append(box)

    # Norte.
    north_x, north_y = 1390, 300
    draw.text((north_x, north_y - 65), "N", fill=BLACK, font=north_font, anchor="mm")
    draw.polygon(
        [
            (north_x, north_y),
            (north_x - 22, north_y + 64),
            (north_x, north_y + 50),
            (north_x + 22, north_y + 64),
        ],
        fill=BLACK,
    )

    output = io.BytesIO()
    image.save(output, format="PNG", optimize=True)
    return output.getvalue()


def find_croqui_media(doc: Document) -> str:
    after_heading = False
    fallback: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().casefold() == "croqui":
            after_heading = True
            continue
        for drawing in paragraph._p.xpath(".//w:drawing"):
            blips = drawing.xpath(".//a:blip")
            if not blips:
                continue
            relationship_id = blips[0].get(qn("r:embed"))
            if relationship_id not in doc.part.rels:
                continue
            target = doc.part.rels[relationship_id].target_ref.replace("\\", "/")
            package_path = target.lstrip("/")
            if not package_path.startswith("word/"):
                package_path = "word/" + package_path
            fallback.append(package_path)
            if after_heading:
                return package_path
    if len(fallback) == 1:
        return fallback[0]
    raise ValueError("Não foi possível identificar a imagem do croqui dentro do DOCX.")


def replace_media(docx_path: Path, output_path: Path, media_name: str, image_data: bytes) -> None:
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
        output_path, "w", compression=zipfile.ZIP_DEFLATED
    ) as target:
        names = set(source.namelist())
        if media_name not in names:
            raise ValueError(f"A mídia do croqui não foi encontrada no pacote: {media_name}")
        for item in source.infolist():
            data = image_data if item.filename == media_name else source.read(item.filename)
            target.writestr(item, data)


def generate_memorial(tabular_path: Path, descriptive_path: Path, output_path: Path) -> dict:
    """Gera o DOCX e retorna informações básicas do resultado."""
    tabular_path = Path(tabular_path).expanduser().resolve()
    descriptive_path = Path(descriptive_path).expanduser().resolve()
    output_path = Path(output_path).expanduser().resolve()
    if not tabular_path.is_file():
        raise FileNotFoundError(f"Memorial tabular não encontrado: {tabular_path}")
    if not descriptive_path.is_file():
        raise FileNotFoundError(f"Memorial descritivo não encontrado: {descriptive_path}")
    if output_path in (tabular_path, descriptive_path):
        raise ValueError(
            "A saída deve ter outro nome. O modelo e o memorial descritivo original são somente leitura."
        )
    output_path.parent.mkdir(parents=True, exist_ok=True)

    memorial = extract_perimeter_memorial(descriptive_path)
    doc = Document(tabular_path)
    update_metadata(doc, memorial)
    table = update_table(doc, memorial)
    croqui_heading = next(
        (paragraph for paragraph in doc.paragraphs if paragraph.text.strip().casefold() == "croqui"),
        None,
    )
    if croqui_heading is not None:
        croqui_heading.paragraph_format.page_break_before = True
        croqui_heading.paragraph_format.keep_with_next = True
        croqui_heading.paragraph_format.keep_together = True
        paragraphs = doc.paragraphs
        heading_index = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is croqui_heading._p)
        for paragraph in paragraphs[heading_index + 1 :]:
            if paragraph._p.xpath(".//w:drawing"):
                break
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
    records = records_from_table(table)
    croqui_media = find_croqui_media(doc)
    croqui_png = render_croqui(records)
    doc.core_properties.modified = datetime.now()

    with tempfile.NamedTemporaryFile(
        prefix="memorial_tabular_", suffix=".docx", dir=output_path.parent, delete=False
    ) as temporary_file:
        temporary_path = Path(temporary_file.name)
    try:
        doc.save(temporary_path)
        replace_media(temporary_path, output_path, croqui_media, croqui_png)
    finally:
        temporary_path.unlink(missing_ok=True)

    with zipfile.ZipFile(output_path) as package:
        if package.testzip() is not None:
            raise ValueError("O DOCX gerado está corrompido.")

    return {"output": output_path, "vertices": len(records)}


def main() -> None:
    args = parse_args()
    tabular_path = args.tabular.expanduser().resolve()
    output_path = (
        args.saida.expanduser().resolve()
        if args.saida
        else tabular_path.with_name(tabular_path.stem + "_Atualizado.docx")
    )
    result = generate_memorial(tabular_path, args.descritivo, output_path)

    print("Memorial gerado com sucesso:")
    print(result["output"])
    print(f"Vértices: {result['vertices']}")


if __name__ == "__main__":
    main()
