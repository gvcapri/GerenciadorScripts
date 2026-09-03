"""Gera o memorial tabular integral de lotes a partir de um DOCX descritivo."""

from __future__ import annotations

import copy
import io
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu

from backend.utils.gerar_memorial_perimetro_tabular import render_croqui, set_cell_text_preserve


VERTEX_RE = re.compile(r"P\s*[-–]?\s*(\d+)", re.IGNORECASE)
COORD_RE = re.compile(
    r"v[ée]rtice\s+(P\s*[-–]?\s*\d+)\s+(?:na|de)\s+coordenada\s*"
    r"\(\s*E\s*X\s*:\s*([\d.]+,\d+)\s+N\s*Y\s*:\s*([\d.]+,\d+)\s*\)",
    re.IGNORECASE,
)
AZIMUTH_RE = re.compile(
    r"azimute\s+(?:de\s+)?(\d{1,3}\s*[°º]\s*\d{1,2}\s*['’]\s*\d{1,2}\s*[\"”]?)",
    re.IGNORECASE,
)


@dataclass
class LotSegment:
    source: str
    destination: str
    easting: str
    northing: str
    confrontation: str
    distance: str
    azimuth: str


@dataclass
class LotMemorial:
    lot: str
    block: str
    area: str
    perimeter: str
    segments: list[LotSegment]


def extract_development(texts: list[str]) -> str:
    for text in texts:
        match = re.search(
            r"LOTEAMENTO\s*:\s*(.+?)(?=\r?\n\s*(?:[ÁA]REA|PER[ÍI]METRO)\s*:|$)",
            text,
            re.IGNORECASE | re.DOTALL,
        )
        if match:
            return clean_space(match.group(1))
    return "LOTEAMENTO"


def normalize_vertex(value: str) -> str:
    match = VERTEX_RE.search(value)
    if not match:
        return value.strip()
    return f"P-{int(match.group(1)):02d}"


def clean_space(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split())


def extract_header(text: str) -> tuple[str, str, str, str] | None:
    identification = re.search(
        r"PROPRIEDADE\s*:\s*LOTE\s*:\s*(.+?)\s*[-–—]\s*QUADRA\s*:\s*([^\r\n]+)",
        text,
        re.IGNORECASE,
    )
    if not identification:
        return None
    area = re.search(r"[ÁA]REA\s*:\s*([\d.,]+)\s*m", text, re.IGNORECASE)
    perimeter = re.search(r"PER[ÍI]METRO\s*:\s*([\d.,]+)\s*m", text, re.IGNORECASE)
    if not area or not perimeter:
        raise ValueError(f"Área ou perímetro ausente no lote {identification.group(1).strip()}.")
    return (
        clean_space(identification.group(1)),
        clean_space(identification.group(2)),
        area.group(1),
        perimeter.group(1),
    )


def extract_segments(description: str, lot_label: str) -> list[LotSegment]:
    coordinates: dict[str, tuple[str, str]] = {}
    coordinate_order: list[str] = []
    for match in COORD_RE.finditer(description):
        vertex = normalize_vertex(match.group(1))
        if vertex not in coordinates:
            coordinate_order.append(vertex)
        coordinates[vertex] = (match.group(2), match.group(3))
    if not coordinate_order:
        raise ValueError(f"Nenhuma coordenada foi encontrada no {lot_label}.")

    azimuths = list(AZIMUTH_RE.finditer(description))
    if not azimuths:
        raise ValueError(f"Nenhum azimute foi encontrado no {lot_label}.")

    current = coordinate_order[0]
    segments: list[LotSegment] = []
    for index, match in enumerate(azimuths):
        end = azimuths[index + 1].start() if index + 1 < len(azimuths) else len(description)
        chunk = description[match.end() : end]
        distance = re.search(
            r"(?:com\s+uma\s+|e\s+)?dist[âa]ncia\s+(?:de\s+)?([\d.,]+)\s*m",
            chunk,
            re.IGNORECASE,
        )
        destination = re.search(r"at[ée]\s+o\s+v[ée]rtice\s+(P\s*[-–]?\s*\d+)", chunk, re.IGNORECASE)
        confrontation = re.search(
            r"confrontando\s+com\s+(.+?)(?=,\s*(?:da[ií]\s+deflete|fechando\b)|[.;]\s*$|$)",
            chunk,
            re.IGNORECASE | re.DOTALL,
        )
        if not distance or not destination or not confrontation:
            raise ValueError(
                f"Não foi possível interpretar o trecho {index + 1} do {lot_label}. "
                "Verifique a redação entre o azimute e o próximo vértice."
            )
        destination_name = normalize_vertex(destination.group(1))
        if current not in coordinates:
            raise ValueError(f"Coordenada ausente para {current} no {lot_label}.")
        easting, northing = coordinates[current]
        segments.append(
            LotSegment(
                source=current,
                destination=destination_name,
                easting=easting,
                northing=northing,
                confrontation=clean_space(confrontation.group(1)).rstrip(" ,;."),
                distance=distance.group(1),
                azimuth=re.sub(r"\s+", "", match.group(1)).replace("º", "°").replace("”", '"'),
            )
        )
        current = destination_name

    if current != segments[0].source:
        raise ValueError(
            f"O perímetro do {lot_label} não fecha: termina em {current}, "
            f"mas começa em {segments[0].source}."
        )
    return segments


def extract_lots(descriptive_path: Path) -> list[LotMemorial]:
    doc = Document(descriptive_path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    header_indexes = [index for index, text in enumerate(paragraphs) if extract_header(text)]
    if not header_indexes:
        raise ValueError("Nenhum cabeçalho de lote (PROPRIEDADE: LOTE / QUADRA) foi encontrado.")

    lots: list[LotMemorial] = []
    for position, header_index in enumerate(header_indexes):
        end = header_indexes[position + 1] if position + 1 < len(header_indexes) else len(paragraphs)
        header = extract_header(paragraphs[header_index])
        assert header is not None
        lot, block, area, perimeter = header
        description = next(
            (
                text
                for text in paragraphs[header_index + 1 : end]
                if re.search(r"Inicia-se\s+a\s+descri[cç][aã]o", text, re.IGNORECASE)
            ),
            "",
        )
        label = f"lote {lot}, quadra {block}"
        if not description:
            raise ValueError(f"Descrição não encontrada para o {label}.")
        lots.append(LotMemorial(lot, block, area, perimeter, extract_segments(description, label)))
    return lots


def rebuild_lot_pages(doc: Document, lot_count: int) -> None:
    """Redimensiona o modelo para qualquer quantidade de páginas de lote."""
    if lot_count < 1:
        raise ValueError("O memorial precisa conter ao menos um lote.")

    body = doc._element.body
    children = list(body.iterchildren())
    first_table_index = next(
        (index for index, child in enumerate(children) if child.tag == qn("w:tbl")),
        None,
    )
    if first_table_index is None or first_table_index + 3 >= len(children):
        raise ValueError("Não foi possível localizar a página-modelo de lote.")

    table_model, title_model, drawing_model, break_model = children[first_table_index : first_table_index + 4]
    if not title_model.xpath(".//w:t[contains(., 'Croqui')]"):
        raise ValueError("O título do croqui não foi encontrado na página-modelo.")
    if not drawing_model.xpath(".//w:drawing"):
        raise ValueError("A imagem do croqui não foi encontrada na página-modelo.")
    if not break_model.xpath(".//w:br[@w:type='page']"):
        raise ValueError("A quebra de página do lote não foi encontrada no modelo.")

    final_section = children[-1]
    if final_section.tag != qn("w:sectPr"):
        raise ValueError("A configuração final de página não foi encontrada no modelo.")
    final_spacer = children[-2]

    for element in children[first_table_index:-1]:
        body.remove(element)

    for index in range(lot_count):
        final_section.addprevious(copy.deepcopy(table_model))
        final_section.addprevious(copy.deepcopy(title_model))
        final_section.addprevious(copy.deepcopy(drawing_model))
        final_section.addprevious(copy.deepcopy(break_model if index < lot_count - 1 else final_spacer))


def update_cover_title(doc: Document, development: str) -> None:
    paragraph = next(
        (
            item
            for item in doc.paragraphs
            if "LISTA INTEGRAL DE MEMORIAIS DESCRITIVOS" in clean_space(item.text).upper()
        ),
        None,
    )
    if paragraph is None:
        return
    text = f"LISTA INTEGRAL DE MEMORIAIS DESCRITIVOS TABULAR DE LOTES DO {development.upper()}"
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    for item in paragraph.runs:
        item.text = ""
    run.text = text


def replace_table_rows(table, lot: LotMemorial) -> None:
    if len(table.rows) < 10 or len(table.columns) < 6:
        raise ValueError("O modelo de lotes contém uma tabela incompatível.")

    set_cell_text_preserve(table.cell(0, 0), f"Quadra {lot.block} - Lote {lot.lot}")
    set_cell_text_preserve(table.cell(6, 0), f"Área: {lot.area} m²")
    set_cell_text_preserve(table.cell(6, 3), f"Perímetro: {lot.perimeter} m")
    set_cell_text_preserve(table.cell(8, 5), "Complemento")

    template_row = copy.deepcopy(table.rows[9]._tr)
    for row in list(table.rows)[9:]:
        table._tbl.remove(row._tr)
    for segment in lot.segments:
        table._tbl.append(copy.deepcopy(template_row))
        row = table.rows[-1]
        values = (
            segment.source,
            segment.easting,
            segment.northing,
            segment.confrontation,
            segment.distance,
            "",
        )
        for column, value in enumerate(values):
            reference = row.cells[4] if column == 5 else None
            set_cell_text_preserve(row.cells[column], value, reference_cell=reference)


def drawing_paragraphs(doc: Document):
    return [paragraph for paragraph in doc.paragraphs if paragraph._p.xpath(".//w:drawing")]


def replace_drawing(paragraph, image_data: bytes, width_emu: int) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    paragraph.add_run().add_picture(io.BytesIO(image_data), width=Emu(width_emu))


def records_for_croqui(lot: LotMemorial) -> list[dict]:
    def number(value: str) -> float:
        return float(value.replace(".", "").replace(",", "."))

    return [
        {
            "de": segment.source,
            "para": segment.destination,
            "e": number(segment.easting),
            "n": number(segment.northing),
            "confrontante": segment.confrontation,
            "distancia": number(segment.distance),
        }
        for segment in lot.segments
    ]


def fitted_croqui_width(base_width: int, segment_count: int) -> int:
    """Reduz croquis de polígonos densos para manter um lote por página."""
    if segment_count >= 8:
        factor = 0.76
    elif segment_count == 7:
        factor = 0.84
    elif segment_count == 6:
        factor = 0.91
    else:
        factor = 1.0
    return int(base_width * factor)


def generate_lot_memorial(template_path: Path, descriptive_path: Path, output_path: Path) -> dict:
    lots = extract_lots(descriptive_path)
    source_doc = Document(descriptive_path)
    development = extract_development([paragraph.text for paragraph in source_doc.paragraphs])
    doc = Document(template_path)
    rebuild_lot_pages(doc, len(lots))
    update_cover_title(doc, development)
    drawings = drawing_paragraphs(doc)
    if len(drawings) != len(lots):
        raise ValueError(
            f"Quantidade divergente de croquis: modelo={len(drawings)}, descritivo={len(lots)}."
        )

    widths = []
    for paragraph in drawings:
        extents = paragraph._p.xpath(".//wp:extent")
        widths.append(int(extents[0].get("cx")) if extents else 4_800_000)

    total_segments = 0
    for index, lot in enumerate(lots):
        replace_table_rows(doc.tables[index], lot)
        croqui = render_croqui(records_for_croqui(lot))
        width = fitted_croqui_width(widths[index], len(lot.segments))
        replace_drawing(drawings[index], croqui, width)
        total_segments += len(lot.segments)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return {
        "output": output_path,
        "type": "lotes",
        "lots": len(lots),
        "segments": total_segments,
    }
