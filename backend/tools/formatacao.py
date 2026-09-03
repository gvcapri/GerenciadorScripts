from pathlib import Path
from typing import Dict, Any, Callable
import os
import re
import unicodedata
from docx import Document
from backend.core.base_tool import BaseTool

class ToolFormatarEspacoMD(BaseTool):
    name = "Formatador de Espaçamento MD"
    description = "Remove parágrafos e Enters vazios de um arquivo Word (.docx)."
    category = "Formatação"
    
    inputs = [
        {"id": "arquivo_md", "label": "Documento DOCX", "type": "document"}
    ]
    
    def execute(self, inputs: Dict[str, Path], output_dir: Path, progress_callback: Callable[[int, str], None], output_filename: str = "") -> Dict[str, Any]:
        caminho = str(inputs["arquivo_md"])
        progress_callback(10, "Lendo documento original...")
        
        doc = Document(caminho)
        total = len(doc.paragraphs)
        removidos = 0
        
        progress_callback(30, "Limpando linhas vazias...")
        
        # Iterar de trás para frente
        for i, p in enumerate(reversed(doc.paragraphs)):
            if i % 10 == 0:
                progress_callback(30 + int((i/total)*60), f"Processando parágrafo {i}/{total}...")
                
            if not p.text.strip() and not p.runs:
                p._element.getparent().remove(p._element)
                removidos += 1
            elif not p.text.strip():
                tem_imagem = False
                for run in p.runs:
                    if run.element.xpath('.//w:drawing') or run.element.xpath('.//w:pict'):
                        tem_imagem = True
                        break
                if not tem_imagem:
                    p._element.getparent().remove(p._element)
                    removidos += 1
                    
        progress_callback(90, "Salvando arquivo limpo...")
        nome_saida = inputs["arquivo_md"].stem + "_LIMPO.docx"
        out_path = self.get_output_filepath(output_dir, nome_saida, output_filename)
        doc.save(str(out_path))
        
        progress_callback(100, f"Finalizado! {removidos} linhas removidas.")
        return {"status": "success", "output_file": str(out_path)}


class ToolMDAutomatico(BaseTool):
    name = "Gerador Automático de MD"
    description = "Converte o arquivo DOCX para PDF e o divide em vários PDFs menores (Ex: QD01_LT_10_01.pdf)."
    category = "Formatação"
    
    inputs = [
        {"id": "arquivo_md", "label": "Documento DOCX", "type": "document"}
    ]
    
    def execute(self, inputs: Dict[str, Path], output_dir: Path, progress_callback: Callable[[int, str], None], output_filename: str = "") -> Dict[str, Any]:
        caminho = str(inputs["arquivo_md"])
        pdf_temp = str(output_dir / "temp_md_automatico.pdf")
        
        try:
            import pythoncom
            pythoncom.CoInitialize()
        except:
            pass
            
        progress_callback(10, "Convertendo DOCX para PDF (Isso abrirá o Word temporariamente)...")
        from docx2pdf import convert
        import sys
        
        # Corrige erro do docx2pdf (sys.stdout = None no PyInstaller windowed)
        class DummyFile:
            def write(self, x): pass
            def flush(self): pass
            
        old_stdout = sys.stdout
        old_stderr = sys.stderr
        if sys.stdout is None: sys.stdout = DummyFile()
        if sys.stderr is None: sys.stderr = DummyFile()

        try:
            convert(caminho, pdf_temp)
        except Exception as e:
            raise RuntimeError(f"Erro ao converter para PDF. Certifique-se de que o Microsoft Word está instalado. Detalhe: {e}")
        finally:
            sys.stdout = old_stdout
            sys.stderr = old_stderr
            
        progress_callback(40, "Abrindo PDF gerado para divisão...")
        from PyPDF2 import PdfReader, PdfWriter
        reader = PdfReader(pdf_temp)
        total_paginas = len(reader.pages)
        
        def pagina_em_branco(texto):
            return not texto or texto.strip() == ""
            
        def normalizar_nome_lote(texto):
            texto = texto.strip()
            if re.fullmatch(r"\d+(?:\s*[\/\-\s]\s*\d+)+", texto):
                numeros = re.findall(r"\d+", texto)
                return "-".join(n.zfill(2) for n in numeros)
            texto = unicodedata.normalize("NFKD", texto)
            texto = texto.encode("ASCII", "ignore").decode("ASCII")
            texto = texto.upper()
            texto = re.sub(r"[^\w\s]", "", texto)
            texto = re.sub(r"\s+", "_", texto)
            return texto.strip("_")
            
        memorial_atual = None
        pagina_contador = 0
        regex_cabecalho = re.compile(r"PROPRIEDADE:\s*LOTE:\s*(.*?)\s*-\s*QUADRA:\s*([0-9A-Z]+)", re.IGNORECASE)
        
        gerados = 0
        
        for i, page in enumerate(reader.pages):
            progress_callback(40 + int((i/total_paginas)*50), f"Processando página {i+1}/{total_paginas}...")
            texto = page.extract_text()
            texto_upper = texto.upper() if texto else ""
            
            if pagina_em_branco(texto_upper):
                continue
                
            match = regex_cabecalho.search(texto_upper)
            if match:
                lote_bruto = match.group(1).strip()
                quadra_bruto = match.group(2).strip()
                quadra = quadra_bruto.zfill(2)
                identificador = normalizar_nome_lote(lote_bruto)
                novo_memorial = (quadra, identificador)
                
                if novo_memorial != memorial_atual:
                    memorial_atual = novo_memorial
                    pagina_contador = 0
                    
            if memorial_atual:
                pagina_contador += 1
                quadra, identificador = memorial_atual
                nome_pdf = f"QD{quadra}_LT_{identificador}_{str(pagina_contador).zfill(2)}.pdf"
                
                writer = PdfWriter()
                writer.add_page(page)
                with open(os.path.join(str(output_dir), nome_pdf), "wb") as f:
                    writer.write(f)
                gerados += 1
                
        # Limpar arquivo temporario
        try:
            if os.path.exists(pdf_temp):
                os.remove(pdf_temp)
        except:
            pass
            
        try:
            import pythoncom
            pythoncom.CoUninitialize()
        except:
            pass
            
        progress_callback(100, f"Divisão finalizada! {gerados} arquivos PDF gerados em {output_dir}")
        return {"status": "success"}

class ToolMemorialPerimetroTabular(BaseTool):
    name = "Gerar Memorial Tabular (Perímetro)"
    description = "Gera um memorial tabular de perímetro com croqui."
    category = "Formatação"
    
    inputs = [
        {"id": "tabular_md", "label": "Modelo Tabular (DOCX)", "type": "document"},
        {"id": "descritivo_md", "label": "Memorial Descritivo (DOCX)", "type": "document"}
    ]
    
    def execute(self, inputs: Dict[str, Path], output_dir: Path, progress_callback: Callable[[int, str], None], output_filename: str = "") -> Dict[str, Any]:
        from backend.utils.gerar_memorial_perimetro_tabular import generate_memorial
        progress_callback(30, "Lendo descritivo e gerando tabular (Isso pode demorar)...")
        tabular_path = inputs["tabular_md"]
        descritivo_path = inputs["descritivo_md"]
        
        default_name = descritivo_path.stem + "_Atualizado.docx"
        saida_path = self.get_output_filepath(output_dir, default_name, output_filename)
        
        generate_memorial(tabular_path, descritivo_path, saida_path)
        progress_callback(100, "Concluído!")
        return {"status": "success", "output_file": str(saida_path)}


class ToolMemorialLotesTabular(BaseTool):
    name = "Gerar Memorial Tabular (Lotes)"
    description = "Gera memoriais tabulares para múltiplos lotes."
    category = "Formatação"
    
    inputs = [
        {"id": "tabular_md", "label": "Modelo Tabular (DOCX)", "type": "document"},
        {"id": "descritivo_md", "label": "Memorial Descritivo (DOCX)", "type": "document"}
    ]
    
    def execute(self, inputs: Dict[str, Path], output_dir: Path, progress_callback: Callable[[int, str], None], output_filename: str = "") -> Dict[str, Any]:
        from backend.utils.gerar_memorial_lotes_tabular import generate_lot_memorial
        progress_callback(30, "Lendo descritivo e gerando tabular (Isso pode demorar)...")
        tabular_path = inputs["tabular_md"]
        descritivo_path = inputs["descritivo_md"]
        
        default_name = descritivo_path.stem + "_Lotes_Tabular.docx"
        saida_path = self.get_output_filepath(output_dir, default_name, output_filename)
        
        res = generate_lot_memorial(tabular_path, descritivo_path, saida_path)
        
        progress_callback(100, f"Concluído! {res.get('lots', 0)} lotes e {res.get('segments', 0)} trechos gerados.")
        return {"status": "success", "output_file": str(saida_path)}


class ToolOrganizarMemoriaisDocx(BaseTool):
    name = "Organizar Memoriais (DOCX)"
    description = "Ordena os memoriais descritivos num arquivo Word por quadra e lote."
    category = "Formatação"
    
    inputs = [
        {"id": "arquivo_md", "label": "Documento DOCX", "type": "document"}
    ]
    
    def execute(self, inputs: Dict[str, Path], output_dir: Path, progress_callback: Callable[[int, str], None], output_filename: str = "") -> Dict[str, Any]:
        from backend.utils.organizar_memoriais_docx import sort_docx
        
        caminho = inputs["arquivo_md"]
        progress_callback(30, "Lendo e organizando memoriais (isso pode demorar)...")
        
        nome_saida = caminho.stem + "_Organizado.docx"
        out_path = self.get_output_filepath(output_dir, nome_saida, output_filename)
        
        total = sort_docx(caminho, out_path)
        
        progress_callback(100, f"Finalizado! {total} memoriais organizados.")
        return {"status": "success", "output_file": str(out_path)}
