from pathlib import Path
from typing import Dict, Any, Callable
import pandas as pd
import re
import unicodedata
from docx import Document
from backend.core.base_tool import BaseTool

class ToolCorrigirCRF(BaseTool):
    name = "Corretor de Arquivos CRF"
    description = "Aplica as correções de Área e Perímetro no DOCX com base em um Gabarito Excel."
    category = "Edições"
    
    inputs = [
        {"id": "excel_gabarito", "label": "Planilha Excel (Gabarito)", "type": "excel"},
        {"id": "docx_crf", "label": "Documento DOCX (CRF)", "type": "document"}
    ]
    
    def execute(self, inputs: Dict[str, Path], output_dir: Path, progress_callback: Callable[[int, str], None], output_filename: str = "") -> Dict[str, Any]:
        caminho_excel = str(inputs["excel_gabarito"])
        caminho_docx = str(inputs["docx_crf"])
        
        progress_callback(10, "Carregando Gabarito Excel...")
        mapa = self.carregar_mapa_excel(caminho_excel, progress_callback)
        if not mapa:
            raise ValueError("Não foi possível carregar o mapa do Excel (Verifique se as colunas Quadra, Lote, Área e Perímetro estão corretas).")
            
        progress_callback(40, "Processando o arquivo Word...")
        nome_saida = inputs["docx_crf"].stem + "_CORRIGIDO.docx"
        caminho_saida = self.get_output_filepath(output_dir, nome_saida, output_filename)
        
        result = self.corrigir_word(caminho_docx, str(caminho_saida), mapa, progress_callback)
        
        progress_callback(100, f"Concluído! Modificações - Áreas: {result['area']} | Perímetros: {result['peri']}")
        return {"status": "success", "output_file": str(caminho_saida)}

    def normalizacao_blindada(self, val):
        if pd.isna(val) or str(val).strip() == "": return ""
        s = str(val).upper().strip()
        s = ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')
        for termo in ["QUADRA", "LOTE", "INDICACAO", "NUMERICA"]:
            s = s.replace(termo, "")
        s = s.strip()
        if "AREA VERDE" in s:
            numero = re.findall(r'\d+', s)
            if numero: return f"AREAVERDE{int(numero[0])}"
            return "AREAVERDE"
        if "AREA REMANESCENTE" in s:
            numero = re.findall(r'\d+', s)
            if numero: return f"AREAREMANESCENTE{int(numero[0])}"
            return "AREAREMANESCENTE"
        s_sem_espaco = re.sub(r'\s+', '', s)
        match_alfanum = re.fullmatch(r'0*(\d+)([A-Z])', s_sem_espaco)
        if match_alfanum:
            return f"{int(match_alfanum.group(1))}{match_alfanum.group(2)}"
        partes = re.split(r'[ \/\-\\&]+', s)
        resultado = [str(int(re.sub(r'[^0-9]', '', p))) for p in partes if re.sub(r'[^0-9]', '', p)]
        if resultado: return "_".join(resultado)
        return re.sub(r'\W+', '', s)

    def float_para_str_br(self, valor):
        try:
            formatado = f"{float(valor):,.2f}"
            return formatado.replace(",", "X").replace(".", ",").replace("X", ".")
        except: return str(valor)

    def carregar_mapa_excel(self, caminho, progress_callback):
        df = pd.read_excel(caminho) 
        colunas_reais = {str(col).strip().upper(): col for col in df.columns}
        col_q = next((orig for limpo, orig in colunas_reais.items() if "QUADRA" in limpo), None)
        col_l = next((orig for limpo, orig in colunas_reais.items() if "LOTE" in limpo), None)
        col_area_alvo = next((orig for limpo, orig in colunas_reais.items() if "AREA" in limpo and ("EXCE" in limpo or "MD" in limpo)), None)
        col_peri_alvo = next((orig for limpo, orig in colunas_reais.items() if "PERI" in limpo and ("EXCE" in limpo or "MD" in limpo)), None)
        if not (col_q and col_l): return {}
        
        mapa = {}
        total = len(df)
        for i, row in df.iterrows():
            if i % 50 == 0: progress_callback(10 + int((i/total)*25), f"Lendo linha {i}/{total} do Excel...")
            q_limpa = self.normalizacao_blindada(row[col_q])
            l_limpo = self.normalizacao_blindada(row[col_l])
            dados = {}
            if col_area_alvo and pd.notna(row[col_area_alvo]):
                try: dados['area'] = float(str(row[col_area_alvo]).replace(",", "."))
                except: pass
            if col_peri_alvo and pd.notna(row[col_peri_alvo]):
                try: dados['peri'] = float(str(row[col_peri_alvo]).replace(",", "."))
                except: pass
            if dados: mapa[(q_limpa, l_limpo)] = dados
        return mapa

    def corrigir_word(self, caminho_entrada, caminho_saida, mapa_correcoes, progress_callback):
        doc = Document(caminho_entrada)
        regex_quadra = re.compile(r"QUADRA\s*([0-9A-Z]+)", re.IGNORECASE)
        regex_lote = re.compile(r"LOTE\s*:?\s*([^\n\r]+)", re.IGNORECASE)
        regex_area = re.compile(r"(ÁREA(?: TOTAL)?\s*:?\s*)([\d\.,]+)(\s*m²?)", re.IGNORECASE)
        regex_peri = re.compile(r"(PERÍMETRO\s*:?\s*)([\d\.,]+)(\s*m?)", re.IGNORECASE)

        quadra_atual, lote_atual = None, None
        count_area, count_peri = 0, 0
        total = len(doc.paragraphs)
        
        for i, p in enumerate(doc.paragraphs):
            if i % 100 == 0: progress_callback(40 + int((i/total)*55), f"Corrigindo Word... {i}/{total}")
            texto = p.text.strip()
            if not texto: continue
            m_q = regex_quadra.search(texto)
            if m_q: quadra_atual = self.normalizacao_blindada(m_q.group(1))
            m_l = regex_lote.search(texto)
            if m_l: lote_atual = self.normalizacao_blindada(m_l.group(1))

            if quadra_atual and lote_atual:
                chave = (quadra_atual, lote_atual)
                if chave in mapa_correcoes:
                    alvos = mapa_correcoes[chave]
                    if "ÁREA" in texto.upper() and 'area' in alvos:
                        m_val = regex_area.search(texto)
                        if m_val:
                            v_antigo, v_novo = m_val.group(2), self.float_para_str_br(alvos['area'])
                            if v_antigo != v_novo:
                                p.text = texto.replace(v_antigo, v_novo)
                                count_area += 1
                                texto = p.text 
                    if "PERÍMETRO" in texto.upper() and 'peri' in alvos:
                        m_val = regex_peri.search(texto)
                        if m_val:
                            v_antigo, v_novo = m_val.group(2), self.float_para_str_br(alvos['peri'])
                            if v_antigo != v_novo:
                                p.text = texto.replace(v_antigo, v_novo)
                                count_peri += 1
        doc.save(caminho_saida)
        return {"area": count_area, "peri": count_peri}
